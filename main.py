import telebot
from telebot import types
import os
import re
import random
import docx2txt
from docx import Document
from settings import API_KEY
from docx.shared import Pt
from docx.oxml.ns import qn

bot = telebot.TeleBot(API_KEY)

# Храним всё по пользователям
user_data = {}

# ---------- КЛАВИАТУРА ----------
def main_keyboard():
    kb = types.ReplyKeyboardMarkup(resize_keyboard=True)
    kb.row("📘 Загрузить компетенции", "🧩 Загрузить вопросы")
    kb.row("🗑 Удалить все файлы")
    kb.row("🧠 Сгенерировать файлы")
    return kb


# ---------- START ----------
@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(
        message.chat.id,
        f"👋 Привет, {message.from_user.first_name or 'пользователь'}!\n\n"
        "Я бот для поиска и генерации файлов по компетенциям 📄\n\n"
        "1️⃣ Загрузите файл с компетенциями (.docx)\n"
        "2️⃣ Загрузите файл с вопросами (.docx)\n"
        "3️⃣ Введите часть названия дисциплины (например: *иностр*, *командн*, *информ*)\n"
        "4️⃣ Нажмите 🧠 *Сгенерировать файлы*\n\n"
        "Я создам Word-файлы только по найденным дисциплинам 📘",
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )


# ---------- ТЕКСТ ----------
@bot.message_handler(content_types=['text'])
def handle_text(message):
    user_id = message.from_user.id
    text = message.text.strip().lower()

    user_dir = f"data_{user_id}"
    comp_file = os.path.join(user_dir, "competencies.docx")
    quest_file = os.path.join(user_dir, "questions.docx")

    os.makedirs(user_dir, exist_ok=True)
    user_data.setdefault(user_id, {})

    # ---- ЗАГРУЗКА ----
    if text == "📘 загрузить компетенции":
        bot.send_message(message.chat.id, "📤 Отправьте Word-файл (.docx) с компетенциями.")
        user_data[user_id]["mode"] = "competencies"
        return

    if text == "🧩 загрузить вопросы":
        bot.send_message(message.chat.id, "📤 Отправьте Word-файл (.docx) с вопросами.")
        user_data[user_id]["mode"] = "questions"
        return

    # ---- УДАЛЕНИЕ ----
    if text == "🗑 удалить все файлы":
        if os.path.exists(user_dir):
            for f in os.listdir(user_dir):
                os.remove(os.path.join(user_dir, f))
            bot.send_message(message.chat.id, "✅ Все файлы удалены.", reply_markup=main_keyboard())
        else:
            bot.send_message(message.chat.id, "⚠️ У вас нет загруженных файлов.", reply_markup=main_keyboard())
        return

    # ---- ГЕНЕРАЦИЯ ----
    if text == "🧠 сгенерировать файлы":
        data = user_data.get(user_id, {})
        found = data.get("found_disciplines")
        if not found:
            bot.send_message(message.chat.id, "⚠️ Сначала введите часть названия дисциплины, чтобы я нашёл нужные.")
            return
        if not os.path.exists(quest_file):
            bot.send_message(message.chat.id, "⚠️ Нужно загрузить файл с вопросами (.docx).")
            return

        bot.send_message(message.chat.id, "⏳ Генерирую файлы, подождите...")

        competencies = data.get("competencies", {})
        questions, _ = extract_questions(quest_file)
        generated = generate_files_per_discipline(user_dir, found, competencies, questions)

        for file_path in generated:
            with open(file_path, "rb") as f:
                bot.send_document(message.chat.id, f)
        bot.send_message(message.chat.id, "✅ Файлы успешно сгенерированы!", reply_markup=main_keyboard())
        return

    # ---- ПОИСК ----
    if not os.path.exists(comp_file):
        bot.send_message(message.chat.id, "⚠️ Сначала загрузите файл с компетенциями (.docx)")
        return

    # Если ещё не извлекали
    if "disciplines" not in user_data[user_id]:
        disciplines = extract_disciplines(comp_file)
        competencies = extract_competencies(comp_file)
        user_data[user_id]["disciplines"] = disciplines
        user_data[user_id]["competencies"] = competencies
        bot.send_message(
            message.chat.id,
            f"✅ Файл загружен! Найдено {len(disciplines)} дисциплин и {len(competencies)} компетенций.\n\n"
            "✏️ Теперь напиши часть названия дисциплины, например: *иностр*, *командн*, *информ*.",
            parse_mode="Markdown"
        )
        return

    disciplines = user_data[user_id]["disciplines"]
    competencies = user_data[user_id]["competencies"]

    found = [d for d in disciplines if text in d.lower()]

    if not found:
        bot.send_message(message.chat.id, "❌ Ничего не найдено. Попробуйте ввести другую часть названия.")
        return

    # --- вывод совпадений ---
    send_long_message(message.chat.id, "📚 Найдено совпадений:\n\n" + "\n\n".join([f"📘 {d}" for d in found]))

    # --- компетенции по найденным дисциплинам ---
    response_lines = []
    for d in found:
        response_lines.append(f"📘 *{d}*")
        # ищем все типы компетенций (УК, ОПК, ПК) — поддержка глубины индексации (напр. 5.3.1)
        comp_codes = re.findall(r"(?:УК|ОПК|ПК)\s*\d+(?:\.\d+)*", d)
        if not comp_codes:
            response_lines.append("⚠️ Нет компетенций для этой дисциплины.\n")
            continue
        for comp in comp_codes:
            comp_key = comp.replace(" ", "")
            desc = find_comp_desc(comp_key, competencies)
            if desc:
                response_lines.append(f"📗 {desc}")
            else:
                response_lines.append(f"⚠️ {comp} — описание не найдено.")
        response_lines.append("")

    send_long_message(
        message.chat.id,
        "📖 *Компетенции, связанные с найденными дисциплинами:*\n\n" + "\n".join(response_lines),
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )

    user_data[user_id]["found_disciplines"] = found


# ---------- ДОКУМЕНТЫ ----------
@bot.message_handler(content_types=['document'])
def handle_document(message):
    user_id = message.from_user.id
    mode = user_data.get(user_id, {}).get("mode")

    if not mode:
        bot.send_message(message.chat.id, "⚠️ Сначала выберите, что загрузить: компетенции или вопросы.")
        return

    user_dir = f"data_{user_id}"
    os.makedirs(user_dir, exist_ok=True)
    file_path = os.path.join(user_dir, f"{mode}.docx")

    file_info = bot.get_file(message.document.file_id)
    downloaded = bot.download_file(file_info.file_path)
    with open(file_path, "wb") as f:
        f.write(downloaded)

    bot.send_message(
        message.chat.id,
        f"✅ Файл '{message.document.file_name}' успешно загружен.",
        reply_markup=main_keyboard()
    )

    comp_file = os.path.join(user_dir, "competencies.docx")
    quest_file = os.path.join(user_dir, "questions.docx")

    # ✅ Если загружены оба файла — только один раз парсим
    if os.path.exists(comp_file) and os.path.exists(quest_file):
        # Проверяем, не были ли уже распознаны
        if "disciplines" not in user_data[user_id] or "competencies" not in user_data[user_id]:
            bot.send_message(message.chat.id, "⏳ Обрабатываю файлы, подождите...")

            disciplines = extract_disciplines(comp_file)
            competencies = extract_competencies(comp_file)

            user_data[user_id]["disciplines"] = disciplines
            user_data[user_id]["competencies"] = competencies

            bot.send_message(
                message.chat.id,
                f"✅ Файлы загружены!\nНайдено {len(disciplines)} дисциплин и {len(competencies)} компетенций.\n\n"
                "✏️ Теперь напиши часть названия дисциплины, например: *иностр*, *командн*, *информ.*",
                parse_mode="Markdown",
                reply_markup=main_keyboard()
            )
        else:
            # Если уже парсили — просто напоминаем, что делать дальше
            bot.send_message(
                message.chat.id,
                "✏️ Теперь напиши часть названия дисциплины, например: *иностр*, *командн*, *информ.*",
                parse_mode="Markdown",
                reply_markup=main_keyboard()
            )


# ---------- ПАРСЕРЫ ----------
def extract_disciplines(file_path):
    full_text = docx2txt.process(file_path)
    print("📘 Текст успешно считан. Общая длина:", len(full_text))
    # Поддерживаем УК, ОПК и ПК в скобках у дисциплины
    pattern = r"(Б\d{1,2}[А-ЯA-Za-zа-яёЁ]*\s*\d*\s*[А-ЯA-Za-zа-яёЁ0-9,\-–\s]+?\((?:УК|ОПК|ПК)\s*[\d.\sА-Яа-яA-ZazlёЁ]*\))"
    matches = re.findall(pattern, full_text)
    print("🔍 Найдено дисциплин:", len(matches))
    disciplines = [" ".join(m.split()) for m in matches]
    return disciplines


def extract_competencies(file_path):
    full_text = docx2txt.process(file_path)
    # Сохраняем переводы строк, но убираем лишние пробелы/табуляции
    cleaned = full_text.replace('\r', '')
    cleaned = re.sub(r"[ \t]+", " ", cleaned).strip()

    # Находим все коды: УК, ОПК, ПК с одной или несколькими точками в номере (напр. 5.3 или 5.3.1 и т.д.)
    code_re = re.compile(r"((?:УК|ОПК|ПК)\s*\d+(?:\.\d+)*)")
    matches = list(code_re.finditer(cleaned))

    # Шаблоны, указывающие на границы блоков, которые не являются описанием компетенции
    stop_patterns = [
        r"\n\s*Б\d",      # следующая дисциплина начинается с Б1...
        r"\n\s*№\s",      # табличная нумерация/заголовок
        r"Код и наименование", r"\bДисциплины\b", r"\bФГОС\b",
        r"\bПС\b", r"\bБ3ГИА\b", r"\bДиректор\b", r"\bЗаведующий\b",
        r"\bПреподаватель\b", r"\bСвязь со стандартами\b"
    ]

    # подстроки для усечения описания
    stop_subs = [
        '\nБ', '\n№', '№ ', 'Код и наименование', 'Дисциплины', 'ФГОС', 'ПС ', 'Б3ГИА',
        'Директор', 'Заведующий', 'Преподаватель', 'Связь со стандартами', 'ПК-', 'УК-', 'ОПК-'
    ]

    competencies = {}
    for i, m in enumerate(matches):
        # Нормализуем найденный код: убираем завершающие точки/запятые/скобки
        code_text_raw = m.group(1)
        code_text = re.sub(r"[\.,;:\)\]]+$", "", code_text_raw).strip()

        start = m.end()
        next_code_start = matches[i + 1].start() if i + 1 < len(matches) else len(cleaned)
        end = next_code_start

        # Ищем ближайший маркер-стоп среди стоп-шаблонов
        for pat in stop_patterns:
            mm = re.search(pat, cleaned[start:next_code_start])
            if mm:
                candidate = start + mm.start()
                if candidate < end:
                    end = candidate

        # Также остановка на двойном переводе строки (новый блок)
        mm = re.search(r"\n\s*\n", cleaned[start:next_code_start])
        if mm:
            candidate = start + mm.start()
            if candidate < end:
                end = candidate

        # Попробуем остановиться на первом конце предложения в пределах разумного (200 символов)
        snippet = cleaned[start:end]
        sent = re.search(r"([\.\!?])\s+", snippet)
        if sent and sent.start() < 200:
            end = start + sent.end()

        desc_raw = cleaned[start:end].strip()

        # Нормализация: убираем ведущие разделители и вкрапления кодов
        desc_raw = re.sub(r"^[\s:;\-–—]+", "", desc_raw)
        desc_raw = re.sub(code_re, "", desc_raw).strip()

        # Усечём по первым стоп-подстрокам, чтобы убрать вкрапления таблиц/заголовков
        earliest = None
        for s in stop_subs:
            idx = desc_raw.find(s)
            if idx != -1:
                if earliest is None or idx < earliest:
                    earliest = idx
        if earliest is not None:
            desc_raw = desc_raw[:earliest].strip()

        # Разбиваем по строкам и убираем строки, которые выглядят как заголовки/номера
        lines = [ln.strip() for ln in desc_raw.splitlines() if ln.strip()]
        clean_lines = []
        for ln in lines:
            if re.match(r"^(?:Б\d|№\s|Код и наименование|Дисциплины|ФГОС|ПС\b|Б3ГИА|Директор|Заведующий|Преподаватель|Связь со стандартами|ПК-|УК-|ОПК-)", ln):
                break
            clean_lines.append(ln)
        desc_raw = ' '.join(clean_lines).strip()

        # Дополнительная усечка по часто встречающимся артефактам (закрывающая скобка + следующий блок, эмодзи и т.п.)
        artifact_patterns = [r"\)\s*Б\d", r"\)\s*Б", r"\)\s*№", r"📘", r"📗", r"⚠️", r"№\s*Код", r"ФГОС", r"ПС\s*\d", r"Б3ГИА"]
        earliest_art = None
        for ap in artifact_patterns:
            a = re.search(ap, desc_raw)
            if a:
                if earliest_art is None or a.start() < earliest_art:
                    earliest_art = a.start()
        if earliest_art is not None:
            desc_raw = desc_raw[:earliest_art].strip()

        # Убираем завершающие служебные символы и одиночные скобки
        desc_raw = re.sub(r"[\-–—\)\(\[\]:;\.,]+$", "", desc_raw).strip()

        # Фолбек: если описание слишком короткое, возьмём чуть более длинный фрагмент до ближайшего логичного конца
        if len(re.sub(r"\s+", "", desc_raw)) < 8:
            extra_end = min(len(cleaned), start + 400)
            candidate_block = cleaned[start:extra_end]
            # обрезаем candidate_block по стоп-паттернам
            for pat in stop_patterns:
                mm = re.search(pat, candidate_block)
                if mm:
                    candidate_block = candidate_block[:mm.start()]
            candidate_block = re.sub(code_re, "", candidate_block).strip()
            # также уберём стоп-подстроки
            earliest2 = None
            for s in stop_subs:
                idx = candidate_block.find(s)
                if idx != -1:
                    if earliest2 is None or idx < earliest2:
                        earliest2 = idx
            if earliest2 is not None:
                candidate_block = candidate_block[:earliest2].strip()
            # и усечём артефакты в candidate_block
            earliest_art2 = None
            for ap in artifact_patterns:
                a = re.search(ap, candidate_block)
                if a:
                    if earliest_art2 is None or a.start() < earliest_art2:
                        earliest_art2 = a.start()
            if earliest_art2 is not None:
                candidate_block = candidate_block[:earliest_art2].strip()
            if len(re.sub(r"\s+", "", candidate_block)) >= 8:
                desc_raw = candidate_block

        # Отбрасываем явно мусорные описания (нет букв)
        if not re.search(r"[А-Яа-яA-Za-z]", desc_raw):
            continue

        # Обрезаем лишнюю длину
        if len(desc_raw) > 400:
            desc_raw = desc_raw[:400].rsplit('.', 1)[0] + "..."

        # Нормализуем ключ (убираем пробелы между префиксом и цифрами)
        key = code_text.replace(" ", "")

        competencies[key] = f"{code_text} — {desc_raw}"

    print("📘 Найдено компетенций:", len(competencies))
    return competencies


def extract_questions(file_path):
    text = docx2txt.process(file_path)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n\n', text)

    sections = [
        "ЕВ", "МВ", "ЧВ", "Соответствие",
        "Одно пропущенное слово", "Два пропущенных слова", "Вложенные вопросы"
    ]

    categorized, current = {}, None
    for line in text.splitlines():
        stripped = line.strip()
        if stripped in sections:
            current = stripped
            categorized[current] = ""
        elif current:
            categorized[current] += line + "\n"

    def normalize_options(options):
        opts = [o.strip() for o in options.splitlines() if o.strip()]
        return "\n".join(opts[:4])

    def find_ev(text):
        matches = re.findall(r"([^\n]+?\?)\s*\n((?:[^\n]*\n){2,8})", text, re.DOTALL)
        return [(q.strip(), normalize_options(o)) for q, o in matches]

    def find_mv(text):
        matches = re.findall(r"([^\n]+?\?)\s*\n((?:[^\n]*\n){2,8})", text, re.DOTALL)
        return [(q.strip(), normalize_options(o)) for q, o in matches]

    def find_chv(text):
        return re.findall(r"([^\n]+?\(Введите[^\n]+?\))\s*\n\s*=\s*([^\n]+)", text, re.DOTALL)

    def find_matching(text):
        blocks = re.findall(r"(Установите соответствие.+?(?=(?:\nУстановите соответствие|$)))", text, re.DOTALL)
        return [re.sub(r'\n{2,}', '\n', b).strip() for b in blocks]

    def find_one_gap(text):
        return re.findall(r"([^\n]+?\(Введите[^\n]+?\))", text)

    def find_two_gap(text):
        blocks = re.split(r'(?=\n?.*?\[\[1\]\].*?\[\[2\]\])', text)
        results = []
        for block in blocks:
            block = block.strip()
            if not block or '[[1]]' not in block:
                continue
            main_part_match = re.search(r'([^\n]*\[\[1\]\].+?\[\[2\]\][^\n]*)', block)
            if not main_part_match:
                continue
            main_part = main_part_match.group(1).strip()
            opt_match = re.search(
                r'(1\s*=\s*[^\n]+(?:\n\s*(?!\d=)[^\n]+)*\n\s*2\s*=\s*[^\n]+(?:\n\s*(?!\[\[)[^\n]+)*)',
                block,
                re.DOTALL
            )
            options = ""
            if opt_match:
                options = "\n" + re.sub(r'\n{2,}', '\n', opt_match.group(1)).strip()
            results.append(f"{main_part}\n{options}".strip())
        return results

    def find_nested(text):
        blocks = re.findall(r"(?:\s*\d+\s*\n)?(.+?(?=\n\s*\d+\s*\n|$))", text, re.DOTALL)
        return [re.sub(r'\n{2,}', '\n', b).strip() for b in blocks if b.strip()]

    extractors = {
        "ЕВ": find_ev, "МВ": find_mv, "ЧВ": find_chv,
        "Соответствие": find_matching,
        "Одно пропущенное слово": find_one_gap,
        "Два пропущенных слова": find_two_gap,
        "Вложенные вопросы": find_nested
    }

    questions = []
    for key, func in extractors.items():
        sec = categorized.get(key, "")
        if not sec.strip():
            continue
        found = func(sec)
        for q in found:
            if isinstance(q, tuple):
                q_text = f"{q[0]}\n{q[1]}"
            else:
                q_text = str(q)
            questions.append(q_text.strip())

    return questions, None


def find_comp_desc(key, competencies):
    """Ищет описание компетенции по ключу.
    Стратегия: точное совпадение -> поиск ключей, начинающихся с key -> поиск по цифровой части -> None
    """
    if key in competencies:
        return competencies[key]

    # Попытка найти более подробные ключи, начинающиеся с данного (например, УК5.3 -> УК5.3.1)
    candidates = [ (k,v) for k,v in competencies.items() if k.startswith(key) or key.startswith(k) ]
    if candidates:
        # выбираем наиболее специфичный (самый длинный ключ)
        best = max(candidates, key=lambda kv: len(kv[0]))
        return best[1]

    # Попытка сопоставления по цифровой части: сравниваем только цифры (например, 53 с 531)
    digits = re.sub(r"\D", "", key)
    if digits:
        for k,v in competencies.items():
            if digits and digits in re.sub(r"\D", "", k):
                return v

    return None


# ---------- ГЕНЕРАЦИЯ ----------
def extract_program_info(file_path):
    """Извлекает направление и профиль из документа компетенций"""
    full_text = docx2txt.process(file_path)
    direction = ""
    profile = ""

    # Ищем строку вида: "по направлению 09.03.01   Информатика и вычислительная техника, профиль - ЭВМ, комплексы, системы и сети"
    match = re.search(
        r"по\s+направлению\s+([\d\.]+\s*[А-Яа-яA-ZazlёЁ\s,]+?)\s*,?\s*профиль\s*[-–—]\s*([А-Яа-яA-ZazlёЁ\s,]+)",
        full_text
    )
    if match:
        direction = match.group(1).strip()
        profile = match.group(2).strip()

    return direction, profile

def generate_files_per_discipline(user_dir, disciplines, competencies, questions):
    generated = []

    comp_file = os.path.join(user_dir, "competencies.docx")
    direction, profile = extract_program_info(comp_file)

    # --- Удаляем возможные вкрапления "Год набора ..." ---
    direction = re.sub(r"год[^\n]*", "", direction, flags=re.IGNORECASE).strip()
    profile = re.sub(r"год[^\n]*", "", profile, flags=re.IGNORECASE).strip()

    if not direction:
        direction = "Направление не указано"
    if not profile:
        profile = "Профиль не указан"

    for disc in disciplines:
        doc = Document()
        question_counter = 1

        # --- Устанавливаем стиль документа (Times New Roman, 14 pt) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # --- Название дисциплины ---
        discipline_match = re.search(r"(Б\d+[А-ЯA-Zazlа-яёЁ0-9\s,\-–]+)", disc)
        discipline_name = discipline_match.group(1).strip() if discipline_match else "Неизвестная дисциплина"

        # --- Ищем коды компетенций ---
        comp_codes = re.findall(r"((?:УК|ОПК|ПК)\s*\d+(?:\.\d+)*)", disc)
        if comp_codes:
            base = re.match(r"((?:УК|ОПК|ПК)\s*\d+)", comp_codes[0])
            short_comp_code = base.group(1).strip() if base else comp_codes[0]
        else:
            short_comp_code = "Компетенция не указана"

        # --- Шапка документа ---
        doc.add_paragraph(f"Задания для компьютерного тестирования по компетенции {short_comp_code}")
        doc.add_paragraph(f"по дисциплине {discipline_name}")
        doc.add_paragraph(f"Направление {direction}")
        doc.add_paragraph(f"Профиль {profile}")
        doc.add_paragraph()

        # --- Таблица с индикаторами ---
        if comp_codes:
            table = doc.add_table(rows=len(comp_codes) + 1, cols=3)
            table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Код компетенции"
            hdr_cells[1].text = "Код индикатора"
            hdr_cells[2].text = "Номера вопросов"

            for i, full_code in enumerate(comp_codes, start=1):
                row = table.rows[i].cells
                base_code = re.match(r"((?:УК|ОПК|ПК)\s*\d+)", full_code).group(1)
                row[0].text = base_code if i == 1 else ""
                row[1].text = full_code.replace(" ", "")
                row[2].text = f"{(i - 1) * 15 + 1}–{i * 15}"

            # Применяем шрифт Times New Roman 14 ко всем ячейкам
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            r.font.size = Pt(14)
        else:
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Код компетенции"
            hdr_cells[1].text = "Код индикатора"
            hdr_cells[2].text = "Номера вопросов"
            table.rows[1].cells[0].text = short_comp_code
            table.rows[1].cells[1].text = f"{short_comp_code}.1"
            table.rows[1].cells[2].text = "1–15"
            table.rows[2].cells[1].text = f"{short_comp_code}.2"
            table.rows[2].cells[2].text = "16–30"

        doc.add_paragraph("\n")

        # --- Основная часть: компетенции и вопросы ---
        for uk in comp_codes:
            uk_key = uk.replace(" ", "")
            desc = find_comp_desc(uk_key, competencies)
            if desc:
                # desc уже в формате 'УК 1.1 — описание'
                desc = re.sub(r"^" + re.escape(uk) + r"\s*[–-]?\s*", "", desc).strip()
                desc = desc.lstrip("—").strip()

                p = doc.add_paragraph()
                run = p.add_run(f"{uk} — {desc}")
                run.bold = True
                p.alignment = 1

                doc.add_paragraph()

                selected = random.sample(questions, min(15, len(questions)))
                for q in selected:
                    doc.add_paragraph(f"{question_counter}. {q}")
                    question_counter += 1

                doc.add_paragraph("\n")

            else:
                p = doc.add_paragraph()
                p.add_run(f"⚠️ {uk} — описание не найдено.")
                p.alignment = 1

        # --- Сохраняем файл ---
        filename = re.sub(r"[^A-Za-zА-Яа-я0-9]", "_", disc[:40]) + ".docx"
        file_path = os.path.join(user_dir, filename)
        doc.save(file_path)
        generated.append(file_path)

    return generated


def send_long_message(chat_id, text, parse_mode=None, reply_markup=None):
    """Отправляет длинный текст частями (безопасно для Telegram)."""
    MAX = 3500
    # Разбиваем по параграфам, чтобы сохранять логические разделы
    paragraphs = text.split('\n\n')
    parts = []
    cur = ''
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        candidate = (cur + '\n\n' + p) if cur else p
        if len(candidate) <= MAX:
            cur = candidate
            continue
        # candidate too big
        if cur:
            parts.append(cur)
            cur = ''
        # если один параграф сам по себе слишком большой — разбиваем по строкам
        if len(p) <= MAX:
            cur = p
        else:
            lines = p.split('\n')
            cur2 = ''
            for ln in lines:
                ln = ln.strip()
                if not ln:
                    continue
                cand2 = (cur2 + '\n' + ln) if cur2 else ln
                if len(cand2) <= MAX:
                    cur2 = cand2
                else:
                    if cur2:
                        parts.append(cur2)
                    # если одна строка длиннее MAX — режем её
                    if len(ln) > MAX:
                        for i in range(0, len(ln), MAX):
                            parts.append(ln[i:i+MAX])
                        cur2 = ''
                    else:
                        cur2 = ln
            if cur2:
                cur = cur2
    if cur:
        parts.append(cur)

    for i, part in enumerate(parts):
        rm = reply_markup if i == len(parts) - 1 else None
        bot.send_message(chat_id, part, parse_mode=parse_mode, reply_markup=rm)


if __name__ == "__main__":
    print("🤖 Бот запущен: поиск и генерация по найденным дисциплинам")
    bot.polling(none_stop=True)
