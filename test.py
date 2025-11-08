import telebot
from telebot import types
import os
import re
import random
import docx2txt
from docx import Document
from settings import API_KEY

bot = telebot.TeleBot(API_KEY)

# Храним всё по пользователям
user_data = {}

# ---------- КЛАВИАТУРА ----------
def main_keyboard():
    kb = types.ReplyKeyboardMarkup(resize_keyboard=True)
    kb.row("📘 Загрузить компетенции", "🧩 Загрузить вопросы")
    kb.row("🗑 Удалить все файлы")
    kb.row("🧠 Сгенерировать файлы")
    return kb


# ---------- START ----------
@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(
        message.chat.id,
        f"👋 Привет, {message.from_user.first_name or 'пользователь'}!\n\n"
        "Я бот для поиска и генерации файлов по компетенциям 📄\n\n"
        "1️⃣ Загрузите файл с компетенциями (.docx)\n"
        "2️⃣ Загрузите файл с вопросами (.docx)\n"
        "3️⃣ Введите часть названия дисциплины (например: *иностр*, *командн*, *информ*)\n"
        "4️⃣ Нажмите 🧠 *Сгенерировать файлы*\n\n"
        "Я создам Word-файлы только по найденным дисциплинам 📘",
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )


# ---------- ТЕКСТ ----------
@bot.message_handler(content_types=['text'])
def handle_text(message):
    user_id = message.from_user.id
    text = message.text.strip().lower()

    user_dir = f"data_{user_id}"
    comp_file = os.path.join(user_dir, "competencies.docx")
    quest_file = os.path.join(user_dir, "questions.docx")

    os.makedirs(user_dir, exist_ok=True)
    user_data.setdefault(user_id, {})

    # ---- ЗАГРУЗКА ----
    if text == "📘 загрузить компетенции":
        bot.send_message(message.chat.id, "📤 Отправьте Word-файл (.docx) с компетенциями.")
        user_data[user_id]["mode"] = "competencies"
        return

    if text == "🧩 загрузить вопросы":
        bot.send_message(message.chat.id, "📤 Отправьте Word-файл (.docx) с вопросами.")
        user_data[user_id]["mode"] = "questions"
        return

    # ---- УДАЛЕНИЕ ----
    if text == "🗑 удалить все файлы":
        if os.path.exists(user_dir):
            for f in os.listdir(user_dir):
                os.remove(os.path.join(user_dir, f))
            bot.send_message(message.chat.id, "✅ Все файлы удалены.", reply_markup=main_keyboard())
        else:
            bot.send_message(message.chat.id, "⚠️ У вас нет загруженных файлов.", reply_markup=main_keyboard())
        return

    # ---- ГЕНЕРАЦИЯ ----
    if text == "🧠 сгенерировать файлы":
        data = user_data.get(user_id, {})
        found = data.get("found_disciplines")
        if not found:
            bot.send_message(message.chat.id, "⚠️ Сначала введите часть названия дисциплины, чтобы я нашёл нужные.")
            return
        if not os.path.exists(quest_file):
            bot.send_message(message.chat.id, "⚠️ Нужно загрузить файл с вопросами (.docx).")
            return

        bot.send_message(message.chat.id, "⏳ Генерирую файлы, подождите...")

        competencies = data.get("competencies", {})
        questions, _ = extract_questions(quest_file)
        generated = generate_files_per_discipline(user_dir, found, competencies, questions)

        for file_path in generated:
            with open(file_path, "rb") as f:
                bot.send_document(message.chat.id, f)
        bot.send_message(message.chat.id, "✅ Файлы успешно сгенерированы!", reply_markup=main_keyboard())
        return

    # ---- ПОИСК ----
    if not os.path.exists(comp_file):
        bot.send_message(message.chat.id, "⚠️ Сначала загрузите файл с компетенциями (.docx)")
        return

    # Если ещё не извлекали
    if "disciplines" not in user_data[user_id]:
        disciplines = extract_disciplines(comp_file)
        competencies = extract_competencies(comp_file)
        user_data[user_id]["disciplines"] = disciplines
        user_data[user_id]["competencies"] = competencies
        bot.send_message(
            message.chat.id,
            f"✅ Файл загружен! Найдено {len(disciplines)} дисциплин и {len(competencies)} компетенций.\n\n"
            "✏️ Теперь напиши часть названия дисциплины, например: *иностр*, *командн*, *информ*.",
            parse_mode="Markdown"
        )
        return

    disciplines = user_data[user_id]["disciplines"]
    competencies = user_data[user_id]["competencies"]

    found = [d for d in disciplines if text in d.lower()]

    if not found:
        bot.send_message(message.chat.id, "❌ Ничего не найдено. Попробуйте ввести другую часть названия.")
        return

    # --- вывод совпадений ---
    bot.send_message(message.chat.id, "📚 Найдено совпадений:\n\n" + "\n\n".join([f"📘 {d}" for d in found]))

    # --- компетенции по найденным дисциплинам ---
    response_lines = []
    for d in found:
        response_lines.append(f"📘 *{d}*")
        uk_codes = re.findall(r"УК\s*\d+\.\d", d)
        if not uk_codes:
            response_lines.append("⚠️ Нет компетенций для этой дисциплины.\n")
            continue
        for uk in uk_codes:
            uk_key = uk.replace(" ", "")
            if uk_key in competencies:
                response_lines.append(f"📗 {competencies[uk_key]}")
            else:
                response_lines.append(f"⚠️ {uk} — описание не найдено.")
        response_lines.append("")

    bot.send_message(
        message.chat.id,
        "📖 *Компетенции, связанные с найденными дисциплинами:*\n\n" + "\n".join(response_lines),
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )

    user_data[user_id]["found_disciplines"] = found


# ---------- ДОКУМЕНТЫ ----------
@bot.message_handler(content_types=['document'])
def handle_document(message):
    user_id = message.from_user.id
    mode = user_data.get(user_id, {}).get("mode")
    if not mode:
        bot.send_message(message.chat.id, "⚠️ Сначала выберите, что загрузить: компетенции или вопросы.")
        return

    user_dir = f"data_{user_id}"
    os.makedirs(user_dir, exist_ok=True)
    file_path = os.path.join(user_dir, f"{mode}.docx")

    file_info = bot.get_file(message.document.file_id)
    downloaded = bot.download_file(file_info.file_path)
    with open(file_path, "wb") as f:
        f.write(downloaded)

    bot.send_message(message.chat.id, f"✅ Файл '{message.document.file_name}' успешно загружен.", reply_markup=main_keyboard())


# ---------- ПАРСЕРЫ ----------
def extract_disciplines(file_path):
    full_text = docx2txt.process(file_path)
    print("📘 Текст успешно считан. Общая длина:", len(full_text))
    pattern = r"(Б\d{1,2}[А-ЯA-Za-zа-яёЁ]*\s*\d*\s*[А-ЯA-Za-zа-яёЁ0-9,\-–\s]+?\(УК\s*[\d.\sА-Яа-яA-Za-z]*\))"
    matches = re.findall(pattern, full_text)
    print("🔍 Найдено дисциплин:", len(matches))
    disciplines = [" ".join(m.split()) for m in matches]
    return disciplines


def extract_competencies(file_path):
    full_text = docx2txt.process(file_path)
    full_text = re.sub(r"\s+", " ", full_text)
    pattern = r"(УК\s*\d+\.\d)(?:\s*[–-]?\s*)([^УБ]+)"
    matches = re.findall(pattern, full_text)
    competencies = {}
    for code, desc in matches:
        clean_code = code.replace(" ", "")
        clean_desc = desc.strip()
        if len(clean_desc) < 10 or "УК" in clean_desc[:10]:
            continue
        if len(clean_desc) > 400:
            clean_desc = clean_desc[:400].rsplit('.', 1)[0] + "..."
        competencies[clean_code] = f"{code} — {clean_desc}"
    print("📘 Найдено компетенций:", len(competencies))
    return competencies


def extract_questions(file_path):
    text = docx2txt.process(file_path)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n\n', text)

    sections = [
        "ЕВ", "МВ", "ЧВ", "Соответствие",
        "Одно пропущенное слово", "Два пропущенных слова", "Вложенные вопросы"
    ]

    categorized, current = {}, None
    for line in text.splitlines():
        stripped = line.strip()
        if stripped in sections:
            current = stripped
            categorized[current] = ""
        elif current:
            categorized[current] += line + "\n"

    def normalize_options(options):
        opts = [o.strip() for o in options.splitlines() if o.strip()]
        return "\n".join(opts[:4])

    def find_ev(text):
        matches = re.findall(r"([^\n]+?\?)\s*\n((?:[^\n]*\n){2,8})", text, re.DOTALL)
        return [(q.strip(), normalize_options(o)) for q, o in matches]

    def find_mv(text):
        matches = re.findall(r"([^\n]+?\?)\s*\n((?:[^\n]*\n){2,8})", text, re.DOTALL)
        return [(q.strip(), normalize_options(o)) for q, o in matches]

    def find_chv(text):
        return re.findall(r"([^\n]+?\(Введите[^\n]+?\))\s*\n\s*=\s*([^\n]+)", text, re.DOTALL)

    def find_matching(text):
        blocks = re.findall(r"(Установите соответствие.+?(?=(?:\nУстановите соответствие|$)))", text, re.DOTALL)
        return [re.sub(r'\n{2,}', '\n', b).strip() for b in blocks]

    def find_one_gap(text):
        return re.findall(r"([^\n]+?\(Введите[^\n]+?\))", text)

    def find_two_gap(text):
        blocks = re.split(r'(?=\n?.*?\[\[1\]\].*?\[\[2\]\])', text)
        results = []
        for block in blocks:
            block = block.strip()
            if not block or '[[1]]' not in block:
                continue
            main_part_match = re.search(r'([^\n]*\[\[1\]\].+?\[\[2\]\][^\n]*)', block)
            if not main_part_match:
                continue
            main_part = main_part_match.group(1).strip()
            opt_match = re.search(
                r'(1\s*=\s*[^\n]+(?:\n\s*(?!\d=)[^\n]+)*\n\s*2\s*=\s*[^\n]+(?:\n\s*(?!\[\[)[^\n]+)*)',
                block,
                re.DOTALL
            )
            options = ""
            if opt_match:
                options = "\n" + re.sub(r'\n{2,}', '\n', opt_match.group(1)).strip()
            results.append(f"{main_part}\n{options}".strip())
        return results

    def find_nested(text):
        blocks = re.findall(r"(?:\s*\d+\s*\n)?(.+?(?=\n\s*\d+\s*\n|$))", text, re.DOTALL)
        return [re.sub(r'\n{2,}', '\n', b).strip() for b in blocks if b.strip()]

    extractors = {
        "ЕВ": find_ev, "МВ": find_mv, "ЧВ": find_chv,
        "Соответствие": find_matching,
        "Одно пропущенное слово": find_one_gap,
        "Два пропущенных слова": find_two_gap,
        "Вложенные вопросы": find_nested
    }

    questions = []
    for key, func in extractors.items():
        sec = categorized.get(key, "")
        if not sec.strip():
            continue
        found = func(sec)
        for q in found:
            if isinstance(q, tuple):
                q_text = f"{q[0]}\n{q[1]}"
            else:
                q_text = str(q)
            questions.append(q_text.strip())

    return questions, None


# ---------- ГЕНЕРАЦИЯ ----------
def generate_files_per_discipline(user_dir, disciplines, competencies, questions):
    generated = []
    for disc in disciplines:
        doc = Document()
        doc.add_heading(disc, level=1)
        uk_codes = re.findall(r"(УК\s*\d+\.\d)", disc)

        for uk in uk_codes:
            uk_key = uk.replace(" ", "")
            if uk_key in competencies:
                doc.add_paragraph(f"📗 {competencies[uk_key]}", style='List Bullet')
                selected = random.sample(questions, min(15, len(questions)))
                for i, q in enumerate(selected, 1):
                    doc.add_paragraph(f"{i}. {q}", style='List Number')

        filename = re.sub(r'[^A-Za-zА-Яа-я0-9]', '_', disc[:40]) + ".docx"
        file_path = os.path.join(user_dir, filename)
        doc.save(file_path)
        generated.append(file_path)

    return generated


# ---------- MAIN ----------
if __name__ == "__main__":
    print("🤖 Бот запущен: поиск и генерация по найденным дисциплинам")
    bot.polling(none_stop=True)
